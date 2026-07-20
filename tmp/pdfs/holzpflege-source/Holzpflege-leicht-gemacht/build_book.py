import argparse
import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Holzpflege-leicht-gemacht.md"
EARTH = RGBColor(91, 70, 54)
INK = RGBColor(44, 42, 38)
MUTED = RGBColor(105, 100, 92)
PAPER = RGBColor(248, 246, 241)

CHAPTERS = [
    "Holzpflege beginnt mit Verstehen",
    "Aufbau, Faser und Maserung",
    "Holzarten erkennen und einordnen",
    "Feuchte, Klima und Holzbewegung",
    "Oberfläche und Zustand bestimmen",
    "Innen oder außen: zwei Belastungswelten",
    "Arbeitsplatz, Werkzeug und Schutz",
    "Schonend reinigen",
    "Probeflächen und Diagnose",
    "Richtig schleifen",
    "Ziehklinge, Hobel und Bürste",
    "Dellen, Risse und lose Verbindungen",
    "Flecken, Ränder und Verfärbungen",
    "Farbe vorbereiten und angleichen",
    "Öle",
    "Hartwachsöle",
    "Wachse",
    "Lacke und Lasuren",
    "Beizen und Farbstoffe",
    "Seife, Lauge und traditionelle Systeme",
    "Schneidebretter und Lebensmittelkontakt",
    "Dielen, Parkett und Treppen",
    "Möbel und Innenausbau",
    "Küche und Bad",
    "Gartenmöbel",
    "Terrassen, Zäune und Sichtschutz",
    "Türen, Fenster und Fassadenholz",
    "Wartungspläne",
    "Fehlerbilder korrigieren",
    "Nachhaltig pflegen",
    "Drei vollständige Praxisprojekte",
    "Schlusskapitel: der ruhige Arbeitsweg",
]

PARTS = [
    ("Teil I", "Holz richtig lesen", range(1, 7)),
    ("Teil II", "Vorbereitung und Reparatur", range(7, 15)),
    ("Teil III", "Oberflächen aufbauen", range(15, 28)),
    ("Teil IV", "Erhalten und sicher entscheiden", range(28, 33)),
]


def set_font(run, name="Calibri", size=None, bold=None, italic=None, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = color


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=8.5, color=MUTED)


def add_markdown_runs(paragraph, text, size=10.5, color=INK):
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=size, bold=True, color=color)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=size, italic=True, color=color)
        else:
            run = paragraph.add_run(token)
            set_font(run, size=size, color=color)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(14.8)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.6)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.22

    specs = {
        "Title": (27, 0, 10, True),
        "Subtitle": (13, 0, 10, False),
        "Heading 1": (16, 18, 8, True),
        "Heading 2": (12.5, 12, 5, True),
        "Heading 3": (11, 9, 4, True),
    }
    for name, (size, before, after, bold) in specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = EARTH
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    styles["Heading 1"].paragraph_format.page_break_before = False

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.font.color.rgb = INK
        style.paragraph_format.left_indent = Cm(0.8)
        style.paragraph_format.first_line_indent = Cm(-0.35)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.18

    if "Figure Caption" not in styles:
        style = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Figure Caption"]
    style.font.name = "Calibri"
    style.font.size = Pt(8.5)
    style.font.italic = True
    style.font.color.rgb = MUTED
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.space_after = Pt(10)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("HOLZPFLEGE LEICHT GEMACHT")
    set_font(hr, size=7.5, bold=True, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(fp, "PAGE")


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("PRAXISBUCH")
    set_font(r, size=9, bold=True, color=EARTH)
    kicker.paragraph_format.space_after = Pt(14)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Holzpflege\nleicht gemacht")
    for run in title.runs:
        set_font(run, size=27, bold=True, color=EARTH)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Holz verstehen, Oberflächen erhalten,\nSchäden sicher beheben")
    set_font(sr, size=13, color=INK)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author.add_run("Thomas Frenzel")
    set_font(ar, size=10.5, bold=True, color=EARTH)
    author.paragraph_format.space_after = Pt(4)

    image = ROOT / "illustrationen" / "01-holzaufbau.png"
    if image.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        inline = run.add_picture(str(image), width=Cm(9.8))
        inline._inline.docPr.set("descr", "Bleistiftzeichnung eines Holzblocks mit sichtbaren Faserflächen")
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(14)

    edition = doc.add_paragraph()
    edition.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = edition.add_run("Erste digitale Ausgabe · 2026")
    set_font(er, size=9.5, color=MUTED)
    doc.add_page_break()


def add_toc(doc, pages):
    h = doc.add_paragraph(style="Heading 1")
    h.add_run("Inhaltsverzeichnis")
    for label, title, nums in PARTS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{label}  {title}")
        set_font(r, size=10, bold=True, color=EARTH)
        for number in nums:
            title_text = CHAPTERS[number - 1]
            entry = doc.add_paragraph()
            entry.paragraph_format.left_indent = Cm(0.25)
            entry.paragraph_format.right_indent = Cm(0.1)
            entry.paragraph_format.space_after = Pt(1.5)
            entry.paragraph_format.line_spacing = 1.0
            tab = entry.paragraph_format.tab_stops
            tab.add_tab_stop(Cm(10.9), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            add_markdown_runs(entry, f"{number}. {title_text}", size=8.8)
            if str(number) in pages:
                pr = entry.add_run(f"\t{pages[str(number)]}")
                set_font(pr, size=8.8, color=MUTED)
    doc.add_page_break()


def add_body(doc, pages):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    body_start = next(i for i, line in enumerate(lines) if line == "# Teil I - Holz richtig lesen")
    front_lines = lines[:body_start]
    skip_toc = False
    for line in front_lines:
        if line == "# Holzpflege leicht gemacht":
            continue
        if line == "# Inhaltsverzeichnis":
            skip_toc = True
            continue
        if skip_toc:
            if line == "---":
                skip_toc = False
            continue
        if line.startswith("## Impressum"):
            p = doc.add_paragraph(style="Heading 1")
            p.add_run("Impressum")
        elif line.startswith("## Sicherheits-"):
            p = doc.add_paragraph(style="Heading 1")
            p.add_run("Sicherheits- und Haftungshinweis")
        elif line == "# Vorwort":
            doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            p.add_run("Vorwort")
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            p.add_run(line[3:])
        elif line and not line.startswith("#") and line != "---" and not line.startswith("Erste digitale Ausgabe") and not line.startswith("## Holz verstehen") and not line.startswith("Praxisbuch für"):
            add_content_line(doc, line)

    add_toc(doc, pages)

    for line in lines[body_start:]:
        if line == "---" or not line:
            continue
        if line.startswith("# Teil "):
            doc.add_page_break()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(55)
            p.paragraph_format.space_after = Pt(8)
            label, title = line[2:].split(" - ", 1)
            r = p.add_run(label.upper())
            set_font(r, size=9, bold=True, color=EARTH)
            t = doc.add_paragraph()
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tr = t.add_run(title)
            set_font(tr, size=22, bold=True, color=EARTH)
            t.paragraph_format.space_after = Pt(28)
        elif re.match(r"# \d+\. ", line):
            p = doc.add_paragraph(style="Heading 1")
            p.add_run(line[2:])
        elif line.startswith("# Anhang"):
            doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            p.add_run(line[2:])
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            p.add_run(line[3:])
        elif line.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            p.add_run(line[2:])
        else:
            add_content_line(doc, line)


def add_content_line(doc, line):
    image_match = re.match(r"!\[(.+?)\]\((.+?)\)", line)
    if image_match:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        run = p.add_run()
        path = ROOT / image_match.group(2)
        inline = run.add_picture(str(path), width=Cm(10.9))
        inline._inline.docPr.set("descr", image_match.group(1))
        return
    if line.startswith("*Abb. "):
        p = doc.add_paragraph(style="Figure Caption")
        p.add_run(line.strip("*"))
        return
    if re.match(r"^- ", line):
        p = doc.add_paragraph(style="List Bullet")
        add_markdown_runs(p, line[2:])
        return
    match = re.match(r"^\d+\. (.+)", line)
    if match:
        p = doc.add_paragraph(style="List Number")
        add_markdown_runs(p, match.group(1))
        return
    p = doc.add_paragraph()
    add_markdown_runs(p, line.replace("  ", " "))
    if line.startswith("Wichtiger Veröffentlichungshinweis:") or line.startswith("Ölgetränkte Lappen"):
        shade_paragraph(p, "F2EEE7")
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.right_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(8)


def split_chapters():
    out = ROOT / "kapitel"
    out.mkdir(exist_ok=True)
    text = SOURCE.read_text(encoding="utf-8")
    matches = list(re.finditer(r"(?m)^# (\d+)\. (.+)$", text))
    for index, match in enumerate(matches):
        start = match.start()
        end = matches[index + 1].start() if index + 1 < len(matches) else text.find("\n# Anhang A", start)
        section = text[start:end].strip() + "\n"
        number = int(match.group(1))
        slug = re.sub(r"[^a-z0-9]+", "-", match.group(2).lower().replace("ä", "ae").replace("ö", "oe").replace("ü", "ue").replace("ß", "ss")).strip("-")
        (out / f"{number:02d}-{slug}.md").write_text(section, encoding="utf-8")


def build(output, pages_file=None):
    pages = {}
    if pages_file and pages_file.exists():
        pages = json.loads(pages_file.read_text(encoding="utf-8"))
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_body(doc, pages)
    props = doc.core_properties
    props.title = "Holzpflege leicht gemacht"
    props.subject = "Praxisbuch über Reinigung, Pflege und Renovierung von Holz"
    props.author = "Thomas Frenzel"
    props.keywords = "Holzpflege, Möbel, Öl, Lack, Außenholz, Reparatur"
    props.comments = "Preset narrative_proposal; named overrides book_handbook_a5 and graphite_earth_palette"
    doc.settings.update_fields_on_open = True
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, default=ROOT / "exports" / "Holzpflege-leicht-gemacht.docx")
    parser.add_argument("--toc-pages", type=Path)
    args = parser.parse_args()
    split_chapters()
    build(args.output, args.toc_pages)


if __name__ == "__main__":
    main()
